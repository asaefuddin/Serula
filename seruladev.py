from flask import Flask, flash, request, redirect, url_for, render_template, make_response, send_from_directory, render_template_string, json
import cv2
import mediapipe as mp
import math
import urllib.request
import os
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from flask import session
from docx.shared import Inches
from io import BytesIO

app = Flask(__name__)

UPLOAD_FOLDER = 'static/upload/'
PROCESSED_FOLDER = 'static/processed/'
RESULTS_FOLDER = 'static/results/'
 
app.secret_key = "secret key"
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['PROCESSED_FOLDER'] = PROCESSED_FOLDER
app.config['RESULTS_FOLDER'] = RESULTS_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
 
ALLOWED_EXTENSIONS = set(['png', 'jpg', 'jpeg', 'gif'])
 

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Inisialisasi modul Mediapipe
mp_drawing = mp.solutions.drawing_utils
mp_pose = mp.solutions.pose
mp_hands = mp.solutions.hands
pose = mp_pose.Pose()
hands = mp_hands.Hands()

# Fungsi untuk menghitung sudut antara tiga titik
def calculate_angle(point1, vertex, point2):
    # Hitung vektor dari point1 ke vertex dan point2 ke vertex
    vector1 = (point1[0] - vertex[0], point1[1] - vertex[1])
    vector2 = (point2[0] - vertex[0], point2[1] - vertex[1])

    # Hitung panjang vektor
    length1 = math.sqrt(vector1[0]**2 + vector1[1]**2)
    length2 = math.sqrt(vector2[0]**2 + vector2[1]**2)

    # Hitung dot product
    dot_product = vector1[0] * vector2[0] + vector1[1] * vector2[1]

    # Hitung cosinus sudut
    cos_angle = dot_product / (length1 * length2)

    # Hitung sudut dalam radian dan konversi ke derajat
    angle_rad = math.acos(cos_angle)
    angle_deg = math.degrees(angle_rad)

    return angle_deg

@app.route('/')
def index():
    filename = session.get('filename')
    if filename is not None and filename != "":
    # Session 'filename' is not empty and contains a value
        print("Session 'filename' is:", filename)
    else:
    # Session 'filename' is empty or not set
        print("Session 'filename' is empty or not set")

    if filename:
        # Jika ada nama file di session, kirim nama file ke template
        return render_template('index.html', css_file="static/styles.css", filename=filename)
    else:
        # Jika tidak ada nama file, kirim None ke template
        return render_template('index.html', css_file="static/styles.css", filename=None)

@app.route('/adjustment', methods=['POST'])
def upload_image():
    if 'file' not in request.files:
        flash('No file part')
        return redirect(request.url)
    
    file = request.files['file']
    
    if file.filename == '':
        flash('No image selected for uploading')
        return redirect(request.url)
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        
        # Set 'filename' in session to the uploaded filename
        session['filename'] = filename
        
        flash('Image successfully uploaded and displayed below')
        
        # Ensure session['filename'] is not empty
        if session.get('filename'):
            filename = session['filename']
        else:
            flash('Session filename is empty or not set.')
            return redirect(url_for('index'))  # Redirect to index if session filename is empty
        
        return render_template('adjustment.html', filename=filename)
    else:
        flash('Allowed image types are - png, jpg, jpeg, gif')
        return redirect(request.url)

@app.route('/adjustment/display/<filename>')
def display_image(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/analyze', methods=['POST'])
def analyze():
    # Ambil gambar dari form HTML
    # Mendapatkan daftar file dalam folder UPLOAD_FOLDER
    files = os.listdir(app.config['UPLOAD_FOLDER'])
    # Mengurutkan file berdasarkan waktu modifikasi (diunggah terakhir)
    latest_file = max(files, key=lambda x: os.path.getmtime(os.path.join(app.config['UPLOAD_FOLDER'], x)))
    # Path lengkap ke gambar yang diunggah terakhir
    image_path = os.path.join(app.config['UPLOAD_FOLDER'], latest_file)

    # Baca gambar dari file
    frame = cv2.imread(image_path)

    # Ubah warna frame menjadi RGB (Mediapipe menggunakan warna RGB)
    rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)

    # Proses pose estimation
    results = pose.process(rgb_frame)

    # Proses deteksi tangan dengan Mediapipe
    hand_results = hands.process(rgb_frame)

    # Recolor back to BGR
    rgb_frame = cv2.cvtColor(rgb_frame, cv2.COLOR_RGB2BGR)

    # Persiapkan data sudut-sudut untuk ditampilkan
    angles_data = {}
 
    # Persiapkan data score untuk ditampilkan
    score_data = {}

    angle_deg_2 = None
    angle_deg_4 = None
    angle_deg_6 = None
    angle_deg_1 = None
    angle_deg_3 = None
    angle_deg_5 = None
    
    # Default values or handle cases where no option was chosen
    upper_arm_position = int(request.form.get('upper_arm_position', 0))
    lower_arm_position = int(request.form.get('lower_arm_position', 0))
    wrist_position = int(request.form.get('wrist_position', 0))
    wrist_twist = int(request.form.get('wrist_twist', 0))
    neck_position = int(request.form.get('neck_position', 0))
    trunk_position = int(request.form.get('trunk_position', 0))
    legs_position = int(request.form.get('legs_position', 0))
    add_muscle_a = int(request.form.get('add_muscle_a', 0))
    add_muscle_b = int(request.form.get('add_muscle_b', 0))
    add_force_a = int(request.form.get('add_force_a', 0))
    add_force_b = int(request.form.get('add_force_b', 0))

    worker_name = request.form.get('worker-name', '')
    job_location = request.form.get('job-location', '')
    job_name = request.form.get('job-name', '')
    image_view = request.form.get('image-view', '')

    # Retrieve the adjustments data
    adjustments_json = request.form.get('adjustments', '{}')
    adjustments = json.loads(adjustments_json)

    # Gambar landmark pose jika hasilnya ditemukan
    if results.pose_landmarks:
        # Memproses landmark jika hasil deteksi tangan tersedia
        for _ in range(30):
            if hand_results.multi_hand_landmarks:
                for hand_landmarks in hand_results.multi_hand_landmarks:
                    # Mendapatkan koordinat Middle Finger MCP
                    middle_finger_mcp = (
                        int(hand_landmarks.landmark[mp_hands.HandLandmark.MIDDLE_FINGER_MCP].x * frame.shape[1]),
                        int(hand_landmarks.landmark[mp_hands.HandLandmark.MIDDLE_FINGER_MCP].y * frame.shape[0])
                    )
                    cv2.circle(frame, middle_finger_mcp, 5, (255, 255, 0), -1)

            # Dapatkan koordinat titik landmark elbow right (--) dan elbow left (--)
            elbow_right = (int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.RIGHT_ELBOW].x * frame.shape[1]),
                        int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.RIGHT_ELBOW].y * frame.shape[0]))
            elbow_left = (int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.LEFT_ELBOW].x * frame.shape[1]),
                        int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.LEFT_ELBOW].y * frame.shape[0]))

            # Dapatkan koordinat titik landmark wrist right (--) dan elbow left (--)
            wrist_right = (int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.RIGHT_WRIST].x * frame.shape[1]),
                        int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.RIGHT_WRIST].y * frame.shape[0]))
            wrist_left = (int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.LEFT_WRIST].x * frame.shape[1]),
                        int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.LEFT_WRIST].y * frame.shape[0]))

            # Dapatkan koordinat titik landmark ear right (--) dan ear left (--)
            ear_right = (int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.RIGHT_EAR].x * frame.shape[1]),
                        int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.RIGHT_EAR].y * frame.shape[0]))
            ear_left = (int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.LEFT_EAR].x * frame.shape[1]),
                        int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.LEFT_EAR].y * frame.shape[0]))

            # Hitung titik tengah (midear) antara ear right dan shoulder left
            midear = ((ear_right[0] + ear_left[0]) // 2, (ear_right[1] + ear_left[1]) // 2)

            # Dapatkan koordinat titik landmark shoulder right (11) dan shoulder left (12)
            shoulder_right = (int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.RIGHT_SHOULDER].x * frame.shape[1]),
                            int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.RIGHT_SHOULDER].y * frame.shape[0]))
            shoulder_left = (int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.LEFT_SHOULDER].x * frame.shape[1]),
                            int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.LEFT_SHOULDER].y * frame.shape[0]))

            # Hitung titik tengah (midsh) antara shoulder right dan shoulder left
            midsh = ((shoulder_right[0] + shoulder_left[0]) // 2, (shoulder_right[1] + shoulder_left[1]) // 2)

            # Dapatkan koordinat titik landmark right Index (19) dan left index (20)
            right_index = (int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.RIGHT_INDEX].x * frame.shape[1]),
                        int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.RIGHT_INDEX].y * frame.shape[0]))
            left_index = (int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.LEFT_INDEX].x * frame.shape[1]),
                        int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.LEFT_INDEX].y * frame.shape[0]))

            # Dapatkan koordinat titik landmark hip right (23) dan hip left (24)
            hip_right = (int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.RIGHT_HIP].x * frame.shape[1]),
                        int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.RIGHT_HIP].y * frame.shape[0]))
            hip_left = (int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.LEFT_HIP].x * frame.shape[1]),
                        int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.LEFT_HIP].y * frame.shape[0]))

            # Hitung titik tengah (midhip) antara hip right dan hip
            midhip = ((hip_right[0] + hip_left[0]) // 2, (hip_right[1] + hip_left[1]) // 2)

            # Dapatkan koordinat titik landmark knee right (25) dan knee left (26)
            knee_right = (int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.RIGHT_KNEE].x * frame.shape[1]),
                        int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.RIGHT_KNEE].y * frame.shape[0]))
            knee_left = (int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.LEFT_KNEE].x * frame.shape[1]),
                        int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.LEFT_KNEE].y * frame.shape[0]))

            # Hitung titik tengah (midkn) antara knee right dan knee left
            midkn = ((knee_right[0] + knee_left[0]) // 2, (knee_right[1] + knee_left[1]) // 2)

            # Dapatkan koordinat titik landmark left foot index (31) dan right foot index (32)
            left_foot_index = (int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.LEFT_FOOT_INDEX].x * frame.shape[1]),
                            int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.LEFT_FOOT_INDEX].y * frame.shape[0]))
            right_foot_index = (int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.RIGHT_FOOT_INDEX].x * frame.shape[1]),
                                int(results.pose_landmarks.landmark[mp_pose.PoseLandmark.RIGHT_FOOT_INDEX].y * frame.shape[0]))

            # Hitung titik tengah (midfoot) antara left foot index dan right foot index
            midfoot = ((left_foot_index[0] + right_foot_index[0]) // 2, (left_foot_index[1] + right_foot_index[1]) // 2)

            # Gambar garis bantu vertikal antara midhip dan midkn (0 derajat)
            vertical_line_x_midhip = midhip[0]  # X-coordinate tetap sama dengan midhip
        
        cv2.line(frame, midhip, (vertical_line_x_midhip, midfoot[1]), (255, 0, 0), 2)

        # Gambar bantu antara midear, midsh, midhip, dan midkn
        cv2.line(frame, midear, midsh, (255, 0, 0), 2)
        cv2.line(frame, midsh, midhip, (255, 0, 0), 2)

        # Gambar titik tengah midear
        cv2.circle(frame, midear, 5, (0, 255, 255), -1)

        # Gambar titik tengah midsh
        cv2.circle(frame, midsh, 5, (0, 255, 255), -1)

        # Gambar titik tengah midhip
        cv2.circle(frame, midhip, 5, (0, 255, 255), -1)

        #Neck
        angle_deg_8 = calculate_angle(midear, midsh, midhip)
        neck_calibrate = angle_deg_8
        #Trunk
        angle_deg_9 = calculate_angle(midsh, midhip, (vertical_line_x_midhip, midfoot[1]))

        # KONDISI DIMULAI DARI SINI
        if image_view == "Right":
            # KANAN
            # Garis bantu antara shoulder,elbow,wrist, dan index kanan
            cv2.line(frame, shoulder_right, elbow_right, (255, 0, 255), 2)
            cv2.line(frame, elbow_right, wrist_right, (255, 0, 255), 2)
            cv2.line(frame, wrist_right, right_index, (255, 0, 255), 2)
            cv2.line(frame, hip_right, shoulder_right, (255, 0, 255), 2)

            # Gambar titik elbow right// Biru
            cv2.circle(frame, elbow_right, 5, (255, 255, 0), -1)

            # Gambar titik shoulder right// Biru
            cv2.circle(frame, shoulder_right, 5, (255, 255, 0), -1)

            # Gambar titik wrist right// Biru
            cv2.circle(frame, wrist_right, 5, (255, 255, 0), -1)

            # Gambar titik right index// Biru
            cv2.circle(frame, right_index, 5, (255, 255, 0), -1)

            # Gambar titik hip right// Biru
            cv2.circle(frame, hip_right, 5, (255, 255, 0), -1)

            # Hitung sudut kanan
            
            angle_deg_2 = calculate_angle(elbow_right, shoulder_right, hip_right)
            angle_deg_4 = calculate_angle(shoulder_right, elbow_right, wrist_right)
            angle_deg_6 = calculate_angle(right_index, wrist_right, elbow_right)

            neck_calibrate = angle_deg_8 
            print(f"Neck Calibrate (initial): {neck_calibrate}")
            
            neck_show = neck_calibrate 
            print(f"Neck Show: {neck_show}")

            neck_calibrate = abs(neck_calibrate) + 20 - 180
            print(f"Neck Calibrate (absolute): {neck_calibrate}")
            
            angle_deg_8 = abs(neck_calibrate)
            print(f"Angle Deg 8: {angle_deg_8}")
                
            # Kondisi Elbow Right Angel
            if 0 <= angle_deg_4 <= 50:
                Lower_Arm_Score = 2
            elif 10 < angle_deg_4 <= 100:
                Lower_Arm_Score = 1
            elif 100 < angle_deg_4 <= 180:
                Lower_Arm_Score = 2
            else:
                Lower_Arm_Score = "N/A" 
            
            #adjustment
            total_lower_arm_score = Lower_Arm_Score + lower_arm_position

            # Kondisi Shoulder Right Angel
            if 0 <= angle_deg_2 <= 20:
                Upper_Arm_Score = 1
            elif 20 < angle_deg_2 <= 45:
                Upper_Arm_Score = 2
            elif 45 < angle_deg_2 <= 90:
                Upper_Arm_Score = 3
            elif 90 < angle_deg_2 <= 180:
                Upper_Arm_Score = 4
            else:
                Upper_Arm_Score = "N/A"
            
            total_upper_arm_score = Upper_Arm_Score + upper_arm_position

            # Kondisi Wrist Right Angle
            Wrist_Score = None
            if angle_deg_6 is not None:
                if 0 <= angle_deg_6 <= 165:
                    Wrist_Score = 3
                elif 165 < angle_deg_6 <= 180:
                    Wrist_Score = 2
                elif angle_deg_6 == 180: 
                    Wrist_Score = 1
                else:
                    Wrist_Score = "N/A" 

            total_wrist_score = Wrist_Score + wrist_position

            # Kondisi Neck Score
            if 170 <= neck_show < 180:
                Neck_Score = 4
            elif 160 <= neck_show < 170:
                Neck_Score = 1
            elif 150 <= neck_show < 160:
                Neck_Score = 2
            else:
                Neck_Score = 3

            total_neck_score = Neck_Score + neck_position

            # Menampilkan sudut-sudut pada layar
            cv2.putText(frame, f'Upper Arm right: {angle_deg_2:.2f} degrees', (10, 30), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 255), 2)
            cv2.putText(frame, f'Lower Arm Right: {angle_deg_4:.2f} degrees', (10, 60), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 255), 2)
            if angle_deg_6 is not None:
                cv2.putText(frame, f'Wrist Right: {angle_deg_6:.2f} degrees', (10, 90), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 255), 2)
            cv2.putText(frame, f'Neck: {neck_calibrate:.2f} degrees', (10, 120), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 255), 2)
            cv2.putText(frame, f'Trunk: {angle_deg_9:.2f} degrees', (10, 150), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 255), 2)
        
        elif image_view == "Left":
            # Garis bantu antara shoulder,elbow,wrist, dan index Kiri
            cv2.line(frame, shoulder_left, elbow_left, (255, 0, 255), 2)
            cv2.line(frame, elbow_left, wrist_left, (255, 0, 255), 2)
            cv2.line(frame, wrist_left, left_index, (255, 0, 255), 2)
            cv2.line(frame, hip_left, shoulder_left, (255, 0, 255), 2)

            # Gambar titik elbow left// Biru
            cv2.circle(frame, elbow_left, 5, (255, 255, 0), -1)

            # Gambar titik shoulder left// Biru
            cv2.circle(frame, shoulder_left, 5, (255, 255, 0), -1)

            # Gambar titik wrist left// Biru
            cv2.circle(frame, wrist_left, 5, (255, 255, 0), -1)

            # Gambar titik left index// Biru
            cv2.circle(frame, left_index, 5, (255, 255, 0), -1)

            # Gambar titik hip left// Biru
            cv2.circle(frame, hip_left, 5, (255, 255, 0), -1)

            # Hitung sudut kiri
            angle_deg_1 = calculate_angle(elbow_left, shoulder_left, hip_left)
            angle_deg_3 = calculate_angle(shoulder_left, elbow_left, wrist_left)
            angle_deg_5 = calculate_angle(left_index, wrist_left, elbow_left)
            
            neck_calibrate = angle_deg_8
            print(f"Neck Calibrate (initial): {neck_calibrate}")
            
            neck_show = neck_calibrate
            print(f"Neck Show: {neck_show}")

            neck_calibrate = abs(neck_calibrate) + 20 - 180
            print(f"Neck Calibrate (absolute): {neck_calibrate}")
            
            angle_deg_8 = abs(neck_calibrate)
            print(f"Angle Deg 8: {angle_deg_8}")

            # Kondisi Elbow Left Angel
            if 0 <= angle_deg_3 <= 50:
                Lower_Arm_Score = 2
            elif 50 < angle_deg_3 <= 100:
                Lower_Arm_Score = 1
            elif 100 < angle_deg_3 <= 180:
                Lower_Arm_Score = 2
            else:
                Lower_Arm_Score = "N/A" 
                
            #adjustment
            total_lower_arm_score = Lower_Arm_Score + lower_arm_position

            # Kondisi Shoulder Left Angel
            if 0 <= angle_deg_1 <= 20:
                Upper_Arm_Score = 1
            elif 20 < angle_deg_1 <= 45:
                Upper_Arm_Score = 2
            elif 45 < angle_deg_1 <= 90:
                Upper_Arm_Score = 3
            elif 90 < angle_deg_1 <= 180:
                Upper_Arm_Score = 4
            else:
                Upper_Arm_Score = "N/A"
                
            total_upper_arm_score = Upper_Arm_Score + upper_arm_position

            # Kondisi Wrist Left Angle
            Wrist_Score = None
            if angle_deg_5 is not None:
                if 0 <= angle_deg_5 <= 165:
                    Wrist_Score = 3
                elif 165 < angle_deg_5 <= 180:
                    Wrist_Score = 2
                elif angle_deg_5 == 180: 
                    Wrist_Score = 1
                else:
                    Wrist_Score = "N/A" 

            total_wrist_score = Wrist_Score + wrist_position
            
            # Kondisi Neck Score
            if 170 <= neck_show < 180:
                Neck_Score = 4
            elif 160 <= neck_show < 170:
                Neck_Score = 1
            elif 150 <= neck_show < 160:
                Neck_Score = 2
            else:
                Neck_Score = 3

            total_neck_score = Neck_Score + neck_position

            # Menampilkan sudut-sudut pada layar
            cv2.putText(frame, f'Upper Arm Left: {angle_deg_1:.2f} degrees', (10, 30), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 255), 2)
            cv2.putText(frame, f'Lower Arm Left: {angle_deg_3:.2f} degrees', (10, 60), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 255), 2)
            if angle_deg_5 is not None:
                cv2.putText(frame, f'Wrist Left: {angle_deg_5:.2f} degrees', (10, 90), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 255), 2)
            cv2.putText(frame, f'Neck: {neck_calibrate:.2f} degrees', (10, 120), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 255), 2)
            cv2.putText(frame, f'Trunk: {angle_deg_9:.2f} degrees', (10, 150), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 255), 2)
        else:
            print("No valid image view selected")
            

        # Kondisi Trunk Score
        if 0 <= angle_deg_9 <= 120:
            Trunk_Score = 4
        elif 120 < angle_deg_9 <= 160:
            Trunk_Score = 3
        elif 160 < angle_deg_9 < 180:
            Trunk_Score = 2
        elif angle_deg_9 == 180:
            Trunk_Score = 1
        else:
            Trunk_Score = "N/A"

        total_trunk_score = Trunk_Score + trunk_position

        Tabel_a = None  # Initialize Tabel_a before using it in conditions
        Tabel_b = None  # Initialize Tabel_b before using it in conditions
        Tabel_c = None  # Initialize Tabel_b before using it in conditions
        final_score = None  # Initialize Tabel_b before using it in conditions
        final_tabel_a = None  # Initialize Tabel_b before using it in conditions
        final_tabel_b = None  # Initialize Tabel_b before using it in conditions

        # Pengelolaan score tabel A
        if total_upper_arm_score == 1 and total_lower_arm_score == 1 and total_wrist_score == 1 and wrist_twist == 1:
            Tabel_a = 1
        #tabel_a=2 total ada 10 kondisi
        elif total_upper_arm_score == 1 and total_lower_arm_score == 2 and total_wrist_score == 1 and wrist_twist == 1: #1211
            Tabel_a = 2
        elif total_upper_arm_score == 1 and total_lower_arm_score == 1 and total_wrist_score == 1 and wrist_twist == 2: #1112
            Tabel_a = 2
        elif total_upper_arm_score == 1 and total_lower_arm_score == 2 and total_wrist_score == 1 and wrist_twist == 2: #1212
            Tabel_a = 2
        elif total_upper_arm_score == 1 and total_lower_arm_score == 3 and total_wrist_score == 1 and wrist_twist == 1: #1311
            Tabel_a = 2
        elif total_upper_arm_score == 2 and total_lower_arm_score == 1 and total_wrist_score == 1 and wrist_twist == 1: #2111
            Tabel_a = 2
        elif total_upper_arm_score == 1 and total_lower_arm_score == 1 and total_wrist_score == 2 and wrist_twist == 1: #1121
            Tabel_a = 2
        elif total_upper_arm_score == 1 and total_lower_arm_score == 2 and total_wrist_score == 2 and wrist_twist == 1: #1221
            Tabel_a = 2
        elif total_upper_arm_score == 1 and total_lower_arm_score == 1 and total_wrist_score == 2 and wrist_twist == 2: #1122
            Tabel_a = 2
        elif total_upper_arm_score == 1 and total_lower_arm_score == 2 and total_wrist_score == 2 and wrist_twist == 2: #1222
            Tabel_a = 2
        elif total_upper_arm_score == 1 and total_lower_arm_score == 1 and total_wrist_score == 3 and wrist_twist == 1: #1131
            Tabel_a = 2
        # table_a=3 total ada 25 kondisi
        elif total_upper_arm_score == 1 and total_lower_arm_score == 1 and total_wrist_score == 3 and wrist_twist == 2: #1132
            Tabel_a = 3
        elif total_upper_arm_score == 1 and total_lower_arm_score == 1 and total_wrist_score == 4 and wrist_twist == 1: #1141
            Tabel_a = 3
        elif total_upper_arm_score == 1 and total_lower_arm_score == 1 and total_wrist_score == 4 and wrist_twist == 2: #1142
            Tabel_a = 3
        elif total_upper_arm_score == 1 and total_lower_arm_score == 2 and total_wrist_score == 3 and wrist_twist == 1: #1231
            Tabel_a = 3
        elif total_upper_arm_score == 1 and total_lower_arm_score == 2 and total_wrist_score == 3 and wrist_twist == 2: #1232
            Tabel_a = 3
        elif total_upper_arm_score == 1 and total_lower_arm_score == 2 and total_wrist_score == 4 and wrist_twist == 1: #1241
            Tabel_a = 3
        elif total_upper_arm_score == 1 and total_lower_arm_score == 2 and total_wrist_score == 4 and wrist_twist == 2: #1242
            Tabel_a = 3
        elif total_upper_arm_score == 1 and total_lower_arm_score == 3 and total_wrist_score == 2 and wrist_twist == 1: #1312
            Tabel_a = 3
        elif total_upper_arm_score == 1 and total_lower_arm_score == 3 and total_wrist_score == 3 and wrist_twist == 1: #1321
            Tabel_a = 3
        elif total_upper_arm_score == 1 and total_lower_arm_score == 3 and total_wrist_score == 3 and wrist_twist == 2: #1322
            Tabel_a = 3
        elif total_upper_arm_score == 1 and total_lower_arm_score == 3 and total_wrist_score == 4 and wrist_twist == 1: #1331
            Tabel_a = 3
        elif total_upper_arm_score == 1 and total_lower_arm_score == 3 and total_wrist_score == 4 and wrist_twist == 2: #1332
            Tabel_a = 3
        elif total_upper_arm_score == 2 and total_lower_arm_score == 1 and total_wrist_score == 2 and wrist_twist == 1: #2112
            Tabel_a = 3
        elif total_upper_arm_score == 2 and total_lower_arm_score == 1 and total_wrist_score == 3 and wrist_twist == 1: #2121
            Tabel_a = 3
        elif total_upper_arm_score == 2 and total_lower_arm_score == 1 and total_wrist_score == 3 and wrist_twist == 2: #2122
            Tabel_a = 3
        elif total_upper_arm_score == 2 and total_lower_arm_score == 1 and total_wrist_score == 4 and wrist_twist == 1: #2131
            Tabel_a = 3
        elif total_upper_arm_score == 2 and total_lower_arm_score == 2 and total_wrist_score == 1 and wrist_twist == 1: #2211
            Tabel_a = 3
        elif total_upper_arm_score == 2 and total_lower_arm_score == 2 and total_wrist_score == 1 and wrist_twist == 2: #2212
            Tabel_a = 3
        elif total_upper_arm_score == 2 and total_lower_arm_score == 2 and total_wrist_score == 2 and wrist_twist == 1: #2221
            Tabel_a = 3
        elif total_upper_arm_score == 2 and total_lower_arm_score == 2 and total_wrist_score == 2 and wrist_twist == 2: #2222
            Tabel_a = 3
        elif total_upper_arm_score == 2 and total_lower_arm_score == 2 and total_wrist_score == 3 and wrist_twist == 1: #2231
            Tabel_a = 3
        elif total_upper_arm_score == 2 and total_lower_arm_score == 3 and total_wrist_score == 1 and wrist_twist == 1: #2311
            Tabel_a = 3
        elif total_upper_arm_score == 3 and total_lower_arm_score == 1 and total_wrist_score == 1 and wrist_twist == 1: #3111
            Tabel_a = 3
        elif total_upper_arm_score == 3 and total_lower_arm_score == 2 and total_wrist_score == 1 and wrist_twist == 1: #3211
            Tabel_a = 3
        elif total_upper_arm_score == 3 and total_lower_arm_score == 1 and total_wrist_score == 1 and wrist_twist == 2: #3112
            Tabel_a = 3
        #tabel_a=4 total ada 40 kondisi
        elif total_upper_arm_score == 1 and total_lower_arm_score == 3 and total_wrist_score == 4 and wrist_twist == 1: #1341
            Tabel_a = 4
        elif total_upper_arm_score == 1 and total_lower_arm_score == 3 and total_wrist_score == 4 and wrist_twist == 2: #1342
            Tabel_a = 4
        elif total_upper_arm_score == 2 and total_lower_arm_score == 1 and total_wrist_score == 3 and wrist_twist == 2: #2132
            Tabel_a = 4
        elif total_upper_arm_score == 2 and total_lower_arm_score == 1 and total_wrist_score == 4 and wrist_twist == 1: #2141
            Tabel_a = 4
        elif total_upper_arm_score == 2 and total_lower_arm_score == 1 and total_wrist_score == 4 and wrist_twist == 2: #2142
            Tabel_a = 4
        elif total_upper_arm_score == 2 and total_lower_arm_score == 2 and total_wrist_score == 3 and wrist_twist == 2: #2232
            Tabel_a = 4
        elif total_upper_arm_score == 2 and total_lower_arm_score == 2 and total_wrist_score == 4 and wrist_twist == 1: #2241
            Tabel_a = 4
        elif total_upper_arm_score == 2 and total_lower_arm_score == 2 and total_wrist_score == 4 and wrist_twist == 2: #2242
            Tabel_a = 4
        elif total_upper_arm_score == 2 and total_lower_arm_score == 3 and total_wrist_score == 1 and wrist_twist == 2: #2312
            Tabel_a = 4
        elif total_upper_arm_score == 2 and total_lower_arm_score == 3 and total_wrist_score == 2 and wrist_twist == 1: #2321
            Tabel_a = 4
        elif total_upper_arm_score == 2 and total_lower_arm_score == 3 and total_wrist_score == 2 and wrist_twist == 2: #2322
            Tabel_a = 4
        elif total_upper_arm_score == 2 and total_lower_arm_score == 3 and total_wrist_score == 3 and wrist_twist == 1: #2331
            Tabel_a = 4
        elif total_upper_arm_score == 2 and total_lower_arm_score == 3 and total_wrist_score == 3 and wrist_twist == 2: #2332
            Tabel_a = 4
        elif total_upper_arm_score == 3 and total_lower_arm_score == 1 and total_wrist_score == 2 and wrist_twist == 1: #3121
            Tabel_a = 4
        elif total_upper_arm_score == 3 and total_lower_arm_score == 1 and total_wrist_score == 2 and wrist_twist == 2: #3122
            Tabel_a = 4
        elif total_upper_arm_score == 3 and total_lower_arm_score == 1 and total_wrist_score == 3 and wrist_twist == 1: #3131
            Tabel_a = 4
        elif total_upper_arm_score == 3 and total_lower_arm_score == 1 and total_wrist_score == 3 and wrist_twist == 2: #3132
            Tabel_a = 4
        elif total_upper_arm_score == 3 and total_lower_arm_score == 2 and total_wrist_score == 1 and wrist_twist == 2: #3212
            Tabel_a = 4
        elif total_upper_arm_score == 3 and total_lower_arm_score == 2 and total_wrist_score == 2 and wrist_twist == 1: #3221
            Tabel_a = 4
        elif total_upper_arm_score == 3 and total_lower_arm_score == 2 and total_wrist_score == 2 and wrist_twist == 2: #3222
            Tabel_a = 4
        elif total_upper_arm_score == 3 and total_lower_arm_score == 2 and total_wrist_score == 3 and wrist_twist == 1: #3231
            Tabel_a = 4
        elif total_upper_arm_score == 3 and total_lower_arm_score == 2 and total_wrist_score == 3 and wrist_twist == 2: #3232
            Tabel_a = 4
        elif total_upper_arm_score == 3 and total_lower_arm_score == 3 and total_wrist_score == 1 and wrist_twist == 1: #3311
            Tabel_a = 4
        elif total_upper_arm_score == 3 and total_lower_arm_score == 3 and total_wrist_score == 2 and wrist_twist == 1: #3312
            Tabel_a = 4
        elif total_upper_arm_score == 3 and total_lower_arm_score == 3 and total_wrist_score == 2 and wrist_twist == 2: #3321
            Tabel_a = 4
        elif total_upper_arm_score == 3 and total_lower_arm_score == 3 and total_wrist_score == 3 and wrist_twist == 1: #3322
            Tabel_a = 4
        elif total_upper_arm_score == 3 and total_lower_arm_score == 3 and total_wrist_score == 3 and wrist_twist == 2: #3331
            Tabel_a = 4
        elif total_upper_arm_score == 4 and total_lower_arm_score == 1 and total_wrist_score == 1 and wrist_twist == 1: #4111
            Tabel_a = 4
        elif total_upper_arm_score == 4 and total_lower_arm_score == 1 and total_wrist_score == 1 and wrist_twist == 2: #4112
            Tabel_a = 4
        elif total_upper_arm_score == 4 and total_lower_arm_score == 1 and total_wrist_score == 2 and wrist_twist == 1: #4121
            Tabel_a = 4
        elif total_upper_arm_score == 4 and total_lower_arm_score == 1 and total_wrist_score == 2 and wrist_twist == 2: #4122
            Tabel_a = 4
        elif total_upper_arm_score == 4 and total_lower_arm_score == 1 and total_wrist_score == 3 and wrist_twist == 1: #4131
            Tabel_a = 4
        elif total_upper_arm_score == 4 and total_lower_arm_score == 2 and total_wrist_score == 1 and wrist_twist == 1: #4211
            Tabel_a = 4
        elif total_upper_arm_score == 4 and total_lower_arm_score == 2 and total_wrist_score == 1 and wrist_twist == 2: #4212
            Tabel_a = 4
        elif total_upper_arm_score == 4 and total_lower_arm_score == 2 and total_wrist_score == 2 and wrist_twist == 1: #4221
            Tabel_a = 4
        elif total_upper_arm_score == 4 and total_lower_arm_score == 2 and total_wrist_score == 2 and wrist_twist == 2: #4222
            Tabel_a = 4
        elif total_upper_arm_score == 4 and total_lower_arm_score == 2 and total_wrist_score == 3 and wrist_twist == 1: #4231
            Tabel_a = 4
        elif total_upper_arm_score == 4 and total_lower_arm_score == 3 and total_wrist_score == 1 and wrist_twist == 1: #4311
            Tabel_a = 4
        elif total_upper_arm_score == 4 and total_lower_arm_score == 3 and total_wrist_score == 1 and wrist_twist == 2: #4312
            Tabel_a = 4
        elif total_upper_arm_score == 4 and total_lower_arm_score == 3 and total_wrist_score == 2 and wrist_twist == 1: #4321
            Tabel_a = 4
        #tabel_a=5 total ada 24 kondisi
        elif total_upper_arm_score == 2 and total_lower_arm_score == 3 and total_wrist_score == 4 and wrist_twist == 1: #2341
            Tabel_a = 5
        elif total_upper_arm_score == 2 and total_lower_arm_score == 3 and total_wrist_score == 4 and wrist_twist == 2: #2342
            Tabel_a = 5
        elif total_upper_arm_score == 3 and total_lower_arm_score == 1 and total_wrist_score == 4 and wrist_twist == 1: #3141
            Tabel_a = 5
        elif total_upper_arm_score == 3 and total_lower_arm_score == 1 and total_wrist_score == 4 and wrist_twist == 2: #3142
            Tabel_a = 5
        elif total_upper_arm_score == 3 and total_lower_arm_score == 2 and total_wrist_score == 4 and wrist_twist == 1: #3241
            Tabel_a = 5
        elif total_upper_arm_score == 3 and total_lower_arm_score == 2 and total_wrist_score == 4 and wrist_twist == 2: #3242
            Tabel_a = 5
        elif total_upper_arm_score == 3 and total_lower_arm_score == 3 and total_wrist_score == 3 and wrist_twist == 2: #3332
            Tabel_a = 5
        elif total_upper_arm_score == 3 and total_lower_arm_score == 3 and total_wrist_score == 4 and wrist_twist == 1: #3341
            Tabel_a = 5
        elif total_upper_arm_score == 3 and total_lower_arm_score == 3 and total_wrist_score == 4 and wrist_twist == 2: #3342
            Tabel_a = 5
        elif total_upper_arm_score == 4 and total_lower_arm_score == 1 and total_wrist_score == 3 and wrist_twist == 2: #4132
            Tabel_a = 5
        elif total_upper_arm_score == 4 and total_lower_arm_score == 1 and total_wrist_score == 4 and wrist_twist == 1: #4141
            Tabel_a = 5
        elif total_upper_arm_score == 4 and total_lower_arm_score == 1 and total_wrist_score == 4 and wrist_twist == 2: #4142
            Tabel_a = 5
        elif total_upper_arm_score == 4 and total_lower_arm_score == 2 and total_wrist_score == 3 and wrist_twist == 2: #4232
            Tabel_a = 5
        elif total_upper_arm_score == 4 and total_lower_arm_score == 2 and total_wrist_score == 4 and wrist_twist == 1: #4241
            Tabel_a = 5
        elif total_upper_arm_score == 4 and total_lower_arm_score == 2 and total_wrist_score == 4 and wrist_twist == 2: #4242
            Tabel_a = 5
        elif total_upper_arm_score == 4 and total_lower_arm_score == 3 and total_wrist_score == 2 and wrist_twist == 2: #4322
            Tabel_a = 5
        elif total_upper_arm_score == 4 and total_lower_arm_score == 3 and total_wrist_score == 3 and wrist_twist == 1: #4331
            Tabel_a = 5
        elif total_upper_arm_score == 4 and total_lower_arm_score == 3 and total_wrist_score == 3 and wrist_twist == 2: #4332
            Tabel_a = 5
        elif total_upper_arm_score == 5 and total_lower_arm_score == 1 and total_wrist_score == 1 and wrist_twist == 1: #5111
            Tabel_a = 5
        elif total_upper_arm_score == 5 and total_lower_arm_score == 1 and total_wrist_score == 1 and wrist_twist == 2: #5112
            Tabel_a = 5
        elif total_upper_arm_score == 5 and total_lower_arm_score == 1 and total_wrist_score == 2 and wrist_twist == 1: #5121
            Tabel_a = 5
        elif total_upper_arm_score == 5 and total_lower_arm_score == 1 and total_wrist_score == 2 and wrist_twist == 2: #5122
            Tabel_a = 5
        elif total_upper_arm_score == 5 and total_lower_arm_score == 1 and total_wrist_score == 3 and wrist_twist == 1: #5131
            Tabel_a = 5
        elif total_upper_arm_score == 5 and total_lower_arm_score == 2 and total_wrist_score == 1 and wrist_twist == 1: #5211
            Tabel_a = 5
        #table_a=6 ada 11 kondisi
        elif total_upper_arm_score == 4 and total_lower_arm_score == 3 and total_wrist_score == 4 and wrist_twist == 1: #4341
            Tabel_a = 6
        elif total_upper_arm_score == 4 and total_lower_arm_score == 3 and total_wrist_score == 4 and wrist_twist == 2: #4342
            Tabel_a = 6
        elif total_upper_arm_score == 5 and total_lower_arm_score == 1 and total_wrist_score == 3 and wrist_twist == 2: #5132
            Tabel_a = 6
        elif total_upper_arm_score == 5 and total_lower_arm_score == 1 and total_wrist_score == 4 and wrist_twist == 1: #5141
            Tabel_a = 6
        elif total_upper_arm_score == 5 and total_lower_arm_score == 2 and total_wrist_score == 1 and wrist_twist == 2: #5212
            Tabel_a = 6
        elif total_upper_arm_score == 5 and total_lower_arm_score == 2 and total_wrist_score == 2 and wrist_twist == 1: #5221
            Tabel_a = 6
        elif total_upper_arm_score == 5 and total_lower_arm_score == 2 and total_wrist_score == 2 and wrist_twist == 2: #5222
            Tabel_a = 6
        elif total_upper_arm_score == 5 and total_lower_arm_score == 2 and total_wrist_score == 3 and wrist_twist == 1: #5231
            Tabel_a = 6
        elif total_upper_arm_score == 5 and total_lower_arm_score == 3 and total_wrist_score == 1 and wrist_twist == 1: #5311
            Tabel_a = 6
        elif total_upper_arm_score == 5 and total_lower_arm_score == 3 and total_wrist_score == 1 and wrist_twist == 2: #5312
            Tabel_a = 6
        elif total_upper_arm_score == 5 and total_lower_arm_score == 3 and total_wrist_score == 2 and wrist_twist == 1: #5321
            Tabel_a = 6
        #table_a=7 ada 13 kondisi
        elif total_upper_arm_score == 5 and total_lower_arm_score == 1 and total_wrist_score == 4 and wrist_twist == 2: #5142
            Tabel_a = 7
        elif total_upper_arm_score == 5 and total_lower_arm_score == 2 and total_wrist_score == 3 and wrist_twist == 2: #5232
            Tabel_a = 7
        elif total_upper_arm_score == 5 and total_lower_arm_score == 2 and total_wrist_score == 4 and wrist_twist == 1: #5241
            Tabel_a = 7
        elif total_upper_arm_score == 5 and total_lower_arm_score == 2 and total_wrist_score == 4 and wrist_twist == 2: #5242
            Tabel_a = 7
        elif total_upper_arm_score == 5 and total_lower_arm_score == 3 and total_wrist_score == 2 and wrist_twist == 2: #5322
            Tabel_a = 7
        elif total_upper_arm_score == 5 and total_lower_arm_score == 3 and total_wrist_score == 3 and wrist_twist == 1: #5331
            Tabel_a = 7
        elif total_upper_arm_score == 5 and total_lower_arm_score == 3 and total_wrist_score == 3 and wrist_twist == 2: #5332
            Tabel_a = 7
        elif total_upper_arm_score == 5 and total_lower_arm_score == 3 and total_wrist_score == 4 and wrist_twist == 1: #5341
            Tabel_a = 7
        elif total_upper_arm_score == 6 and total_lower_arm_score == 1 and total_wrist_score == 1 and wrist_twist == 1: #6111
            Tabel_a = 7
        elif total_upper_arm_score == 6 and total_lower_arm_score == 1 and total_wrist_score == 1 and wrist_twist == 2: #6112
            Tabel_a = 7
        elif total_upper_arm_score == 6 and total_lower_arm_score == 1 and total_wrist_score == 2 and wrist_twist == 1: #6121
            Tabel_a = 7
        elif total_upper_arm_score == 6 and total_lower_arm_score == 1 and total_wrist_score == 2 and wrist_twist == 2: #6122
            Tabel_a = 7
        elif total_upper_arm_score == 6 and total_lower_arm_score == 1 and total_wrist_score == 3 and wrist_twist == 1: #6131
            Tabel_a = 7
        #table_a=8 ada 8 kondisi
        elif total_upper_arm_score == 5 and total_lower_arm_score == 3 and total_wrist_score == 4 and wrist_twist == 2: #5342
            Tabel_a = 8
        elif total_upper_arm_score == 6 and total_lower_arm_score == 1 and total_wrist_score == 3 and wrist_twist == 2: #6132
            Tabel_a = 8
        elif total_upper_arm_score == 6 and total_lower_arm_score == 1 and total_wrist_score == 4 and wrist_twist == 1: #6141
            Tabel_a = 8
        elif total_upper_arm_score == 6 and total_lower_arm_score == 2 and total_wrist_score == 1 and wrist_twist == 1: #6211
            Tabel_a = 8
        elif total_upper_arm_score == 6 and total_lower_arm_score == 2 and total_wrist_score == 1 and wrist_twist == 2: #6212
            Tabel_a = 8
        elif total_upper_arm_score == 6 and total_lower_arm_score == 2 and total_wrist_score == 2 and wrist_twist == 1: #6221
            Tabel_a = 8
        elif total_upper_arm_score == 6 and total_lower_arm_score == 2 and total_wrist_score == 2 and wrist_twist == 2: #6222
            Tabel_a = 8
        elif total_upper_arm_score == 6 and total_lower_arm_score == 2 and total_wrist_score == 3 and wrist_twist == 1: #6231
            Tabel_a = 8
        #table_a=9 ada 12 kondisi
        elif total_upper_arm_score == 6 and total_lower_arm_score == 1 and total_wrist_score == 4 and wrist_twist == 2: #6142
            Tabel_a = 9
        elif total_upper_arm_score == 6 and total_lower_arm_score == 2 and total_wrist_score == 3 and wrist_twist == 2: #6232
            Tabel_a = 9
        elif total_upper_arm_score == 6 and total_lower_arm_score == 2 and total_wrist_score == 4 and wrist_twist == 1: #6241
            Tabel_a = 9
        elif total_upper_arm_score == 6 and total_lower_arm_score == 2 and total_wrist_score == 4 and wrist_twist == 2: #6242
            Tabel_a = 9
        elif total_upper_arm_score == 6 and total_lower_arm_score == 3 and total_wrist_score == 1 and wrist_twist == 1: #6311
            Tabel_a = 9
        elif total_upper_arm_score == 6 and total_lower_arm_score == 3 and total_wrist_score == 1 and wrist_twist == 2: #6312
            Tabel_a = 9
        elif total_upper_arm_score == 6 and total_lower_arm_score == 3 and total_wrist_score == 2 and wrist_twist == 1: #6321
            Tabel_a = 9
        elif total_upper_arm_score == 6 and total_lower_arm_score == 3 and total_wrist_score == 2 and wrist_twist == 2: #6322
            Tabel_a = 9
        elif total_upper_arm_score == 6 and total_lower_arm_score == 3 and total_wrist_score == 3 and wrist_twist == 1: #6331
            Tabel_a = 9
        elif total_upper_arm_score == 6 and total_lower_arm_score == 3 and total_wrist_score == 3 and wrist_twist == 2: #6332
            Tabel_a = 9
        elif total_upper_arm_score == 6 and total_lower_arm_score == 3 and total_wrist_score == 4 and wrist_twist == 1: #6341
            Tabel_a = 9
        elif total_upper_arm_score == 6 and total_lower_arm_score == 3 and total_wrist_score == 4 and wrist_twist == 2: #6342
            Tabel_a = 9
        else:
            Tabel_a = 0

        final_tabel_a = Tabel_a + add_muscle_a + add_force_a

        #table_b
        if total_neck_score == 1 and total_trunk_score == 1 and legs_position == 1: #111
            Tabel_b = 1
        #tabel_b=2 ada 3 kondisi
        elif total_neck_score == 1 and total_trunk_score == 2 and legs_position == 1: #121
            Tabel_b = 2
        elif total_neck_score == 2 and total_trunk_score == 1 and legs_position == 1: #211
            Tabel_b = 2
        elif total_neck_score == 2 and total_trunk_score == 2 and legs_position == 1: #221
            Tabel_b = 2
        #tabel_b=3 ada 8 kondisi
        elif total_neck_score == 1 and total_trunk_score == 1 and legs_position == 2: #112
            Tabel_b = 3
        elif total_neck_score == 1 and total_trunk_score == 2 and legs_position == 2: #122
            Tabel_b = 3
        elif total_neck_score == 1 and total_trunk_score == 3 and legs_position == 1: #131
            Tabel_b = 3
        elif total_neck_score == 2 and total_trunk_score == 1 and legs_position == 2: #212
            Tabel_b = 3
        elif total_neck_score == 2 and total_trunk_score == 2 and legs_position == 2: #222
            Tabel_b = 3
        elif total_neck_score == 3 and total_trunk_score == 1 and legs_position == 1: #311
            Tabel_b = 3
        elif total_neck_score == 3 and total_trunk_score == 1 and legs_position == 2: #312
            Tabel_b = 3
        elif total_neck_score == 3 and total_trunk_score == 2 and legs_position == 1: #321
            Tabel_b = 3
        #tabel_b=4 ada 4 kondisi
        elif total_neck_score == 1 and total_trunk_score == 3 and legs_position == 2: #132
            Tabel_b = 4
        elif total_neck_score == 2 and total_trunk_score == 3 and legs_position == 1: #231
            Tabel_b = 4
        elif total_neck_score == 3 and total_trunk_score == 2 and legs_position == 2: #322
            Tabel_b = 4
        elif total_neck_score == 3 and total_trunk_score == 3 and legs_position == 1: #331
            Tabel_b = 4
        #tabel_b=5 ada 10 kondisi
        elif total_neck_score == 1 and total_trunk_score == 4 and legs_position == 1: #141
            Tabel_b = 5
        elif total_neck_score == 1 and total_trunk_score == 4 and legs_position == 2: #142
            Tabel_b = 5
        elif total_neck_score == 2 and total_trunk_score == 3 and legs_position == 2: #232
            Tabel_b = 5
        elif total_neck_score == 2 and total_trunk_score == 4 and legs_position == 1: #241
            Tabel_b = 5
        elif total_neck_score == 2 and total_trunk_score == 4 and legs_position == 2: #242
            Tabel_b = 5
        elif total_neck_score == 3 and total_trunk_score == 3 and legs_position == 2: #332
            Tabel_b = 5
        elif total_neck_score == 3 and total_trunk_score == 4 and legs_position == 1: #341
            Tabel_b = 5
        elif total_neck_score == 4 and total_trunk_score == 1 and legs_position == 1: #411
            Tabel_b = 5
        elif total_neck_score == 4 and total_trunk_score == 1 and legs_position == 2: #412
            Tabel_b = 5
        elif total_neck_score == 4 and total_trunk_score == 2 and legs_position == 1: #421
            Tabel_b = 5
        #tabel_b=6 ada 7 kondisi
        elif total_neck_score == 1 and total_trunk_score == 5 and legs_position == 1: #151
            Tabel_b = 6
        elif total_neck_score == 1 and total_trunk_score == 5 and legs_position == 2: #152
            Tabel_b = 6
        elif total_neck_score == 2 and total_trunk_score == 5 and legs_position == 1: #251
            Tabel_b = 6
        elif total_neck_score == 3 and total_trunk_score == 4 and legs_position == 2: #342
            Tabel_b = 6
        elif total_neck_score == 3 and total_trunk_score == 5 and legs_position == 1: #351
            Tabel_b = 6
        elif total_neck_score == 4 and total_trunk_score == 2 and legs_position == 2: #422
            Tabel_b = 6
        elif total_neck_score == 4 and total_trunk_score == 3 and legs_position == 1: #431
            Tabel_b = 6
        #tabel_b=7 ada 18 kondisi
        elif total_neck_score == 1 and total_trunk_score == 6 and legs_position == 1: #161
            Tabel_b = 7
        elif total_neck_score == 1 and total_trunk_score == 6 and legs_position == 2: #162
            Tabel_b = 7
        elif total_neck_score == 2 and total_trunk_score == 5 and legs_position == 2: #252
            Tabel_b = 7
        elif total_neck_score == 2 and total_trunk_score == 6 and legs_position == 1: #261
            Tabel_b = 7
        elif total_neck_score == 2 and total_trunk_score == 6 and legs_position == 2: #262
            Tabel_b = 7
        elif total_neck_score == 3 and total_trunk_score == 5 and legs_position == 2: #352
            Tabel_b = 7
        elif total_neck_score == 3 and total_trunk_score == 6 and legs_position == 1: #361
            Tabel_b = 7
        elif total_neck_score == 3 and total_trunk_score == 6 and legs_position == 2: #362
            Tabel_b = 7
        elif total_neck_score == 4 and total_trunk_score == 3 and legs_position == 2: #432
            Tabel_b = 7
        elif total_neck_score == 4 and total_trunk_score == 4 and legs_position == 1: #441
            Tabel_b = 7
        elif total_neck_score == 4 and total_trunk_score == 4 and legs_position == 2: #442
            Tabel_b = 7
        elif total_neck_score == 4 and total_trunk_score == 5 and legs_position == 1: #451
            Tabel_b = 7
        elif total_neck_score == 4 and total_trunk_score == 5 and legs_position == 2: #452
            Tabel_b = 7
        elif total_neck_score == 5 and total_trunk_score == 1 and legs_position == 1: #511
            Tabel_b = 7
        elif total_neck_score == 5 and total_trunk_score == 1 and legs_position == 2: #512
            Tabel_b = 7
        elif total_neck_score == 5 and total_trunk_score == 2 and legs_position == 1: #521
            Tabel_b = 7
        elif total_neck_score == 5 and total_trunk_score == 2 and legs_position == 2: #522
            Tabel_b = 7
        elif total_neck_score == 5 and total_trunk_score == 3 and legs_position == 1: #531
            Tabel_b = 7
        #tabel_b=8 ada 16 kondisi
        elif total_neck_score == 4 and total_trunk_score == 6 and legs_position == 1: #461
            Tabel_b = 8
        elif total_neck_score == 4 and total_trunk_score == 6 and legs_position == 2: #462
            Tabel_b = 8
        elif total_neck_score == 5 and total_trunk_score == 3 and legs_position == 2: #532
            Tabel_b = 8
        elif total_neck_score == 5 and total_trunk_score == 4 and legs_position == 1: #541
            Tabel_b = 8
        elif total_neck_score == 5 and total_trunk_score == 4 and legs_position == 2: #542
            Tabel_b = 8
        elif total_neck_score == 5 and total_trunk_score == 5 and legs_position == 1: #551
            Tabel_b = 8
        elif total_neck_score == 5 and total_trunk_score == 5 and legs_position == 2: #552
            Tabel_b = 8
        elif total_neck_score == 5 and total_trunk_score == 6 and legs_position == 1: #561
            Tabel_b = 8
        elif total_neck_score == 5 and total_trunk_score == 6 and legs_position == 2: #562
            Tabel_b = 8
        elif total_neck_score == 6 and total_trunk_score == 1 and legs_position == 1: #611
            Tabel_b = 8
        elif total_neck_score == 6 and total_trunk_score == 1 and legs_position == 2: #612
            Tabel_b = 8
        elif total_neck_score == 6 and total_trunk_score == 2 and legs_position == 1: #621
            Tabel_b = 8
        elif total_neck_score == 6 and total_trunk_score == 2 and legs_position == 2: #622
            Tabel_b = 8
        elif total_neck_score == 6 and total_trunk_score == 3 and legs_position == 1: #631
            Tabel_b = 8
        elif total_neck_score == 6 and total_trunk_score == 3 and legs_position == 2: #632
            Tabel_b = 8
        elif total_neck_score == 6 and total_trunk_score == 4 and legs_position == 1: #641
            Tabel_b = 8
        elif total_neck_score == 6 and total_trunk_score == 4 and legs_position == 2: #642
            Tabel_b = 9
        elif total_neck_score == 6 and total_trunk_score == 5 and legs_position == 1: #651
            Tabel_b = 9
        elif total_neck_score == 6 and total_trunk_score == 5 and legs_position == 2: #652
            Tabel_b = 9
        elif total_neck_score == 6 and total_trunk_score == 6 and legs_position == 1: #661
            Tabel_b = 9
        elif total_neck_score == 6 and total_trunk_score == 6 and legs_position == 2: #662
            Tabel_b = 9
        else:
            Tabel_b = 0

        final_tabel_b = Tabel_b + add_muscle_b + add_force_b

        # Tabel C (Matriks Tabel A dan Tabel B)
        if final_tabel_a == 1 and final_tabel_b == 1:  # 11
            Tabel_c = 1
        elif final_tabel_a == 1 and final_tabel_b == 2:  # 12
            Tabel_c = 2
        elif final_tabel_a == 2 and final_tabel_b == 1:  # 21
            Tabel_c = 2
        elif final_tabel_a == 2 and final_tabel_b == 2:  # 22
            Tabel_c = 2
        elif final_tabel_a == 1 and final_tabel_b == 3:  # 13
            Tabel_c = 3
        elif final_tabel_a == 1 and final_tabel_b == 4:  # 14
            Tabel_c = 3
        elif final_tabel_a == 2 and final_tabel_b == 3:  # 23
            Tabel_c = 3
        elif final_tabel_a == 3 and final_tabel_b == 1:  # 31
            Tabel_c = 3
        elif final_tabel_a == 3 and final_tabel_b == 2:  # 32
            Tabel_c = 3
        elif final_tabel_a == 3 and final_tabel_b == 3:  # 33
            Tabel_c = 3
        elif final_tabel_a == 4 and final_tabel_b == 1:  # 41
            Tabel_c = 3
        elif final_tabel_a == 4 and final_tabel_b == 2:  # 42
            Tabel_c = 3
        elif final_tabel_a == 4 and final_tabel_b == 3:  # 43
            Tabel_c = 3
        elif final_tabel_a == 1 and final_tabel_b == 5:  # 15
            Tabel_c = 4
        elif final_tabel_a == 2 and final_tabel_b == 4:  # 24
            Tabel_c = 4
        elif final_tabel_a == 2 and final_tabel_b == 5:  # 25
            Tabel_c = 4
        elif final_tabel_a == 3 and final_tabel_b == 4:  # 34
            Tabel_c = 4
        elif final_tabel_a == 3 and final_tabel_b == 5:  # 35
            Tabel_c = 4
        elif final_tabel_a == 4 and final_tabel_b == 4:  # 44
            Tabel_c = 4
        elif final_tabel_a == 5 and final_tabel_b == 1:  # 51
            Tabel_c = 4
        elif final_tabel_a == 5 and final_tabel_b == 2:  # 52
            Tabel_c = 4
        elif final_tabel_a == 5 and final_tabel_b == 3:  # 53
            Tabel_c = 4
        elif final_tabel_a == 6 and final_tabel_b == 1:  # 61
            Tabel_c = 4
        elif final_tabel_a == 6 and final_tabel_b == 2:  # 62
            Tabel_c = 4
        elif final_tabel_a == 1 and final_tabel_b == 6:  # 16
            Tabel_c = 5
        elif final_tabel_a == 1 and final_tabel_b >= 7:  # 17
            Tabel_c = 5
        elif final_tabel_a == 2 and final_tabel_b == 6:  # 26
            Tabel_c = 5
        elif final_tabel_a == 2 and final_tabel_b >= 7:  # 27
            Tabel_c = 5
        elif final_tabel_a == 3 and final_tabel_b == 6:  # 36
            Tabel_c = 5
        elif final_tabel_a == 4 and final_tabel_b == 5:  # 45
            Tabel_c = 5
        elif final_tabel_a == 5 and final_tabel_b == 4:  # 54
            Tabel_c = 5
        elif final_tabel_a == 6 and final_tabel_b == 3:  # 63
            Tabel_c = 5
        elif final_tabel_a == 7 and final_tabel_b == 1:  # 71
            Tabel_c = 5
        elif final_tabel_a == 7 and final_tabel_b == 2:  # 72
            Tabel_c = 5
        elif final_tabel_a >= 8 and final_tabel_b == 1:  # 81
            Tabel_c = 5
        elif final_tabel_a >= 8 and final_tabel_b == 2:  # 82
            Tabel_c = 5
        elif final_tabel_a == 3 and final_tabel_b >= 7:  # 37
            Tabel_c = 6
        elif final_tabel_a == 4 and final_tabel_b == 6:  # 46
            Tabel_c = 6
        elif final_tabel_a == 4 and final_tabel_b >= 7:  # 47
            Tabel_c = 6
        elif final_tabel_a == 5 and final_tabel_b == 5:  # 55
            Tabel_c = 6
        elif final_tabel_a == 6 and final_tabel_b == 4:  # 64
            Tabel_c = 6
        elif final_tabel_a == 6 and final_tabel_b == 5:  # 65
            Tabel_c = 6
        elif final_tabel_a == 7 and final_tabel_b == 3:  # 73
            Tabel_c = 6
        elif final_tabel_a == 7 and final_tabel_b == 4:  # 74
            Tabel_c = 6
        elif final_tabel_a >= 8 and final_tabel_b == 3:  # 83
            Tabel_c = 6
        elif final_tabel_a == 5 and final_tabel_b == 6:  # 56
            Tabel_c = 7
        elif final_tabel_a == 5 and final_tabel_b >= 7:  # 57
            Tabel_c = 7
        elif final_tabel_a == 6 and final_tabel_b == 6:  # 66
            Tabel_c = 7
        elif final_tabel_a == 6 and final_tabel_b >= 7:  # 67
            Tabel_c = 7
        elif final_tabel_a == 7 and final_tabel_b == 5:  # 75
            Tabel_c = 7
        elif final_tabel_a == 7 and final_tabel_b == 6:  # 76
            Tabel_c = 7
        elif final_tabel_a == 7 and final_tabel_b >= 7:  # 77
            Tabel_c = 7
        elif final_tabel_a >= 8 and final_tabel_b == 4:  # 84
            Tabel_c = 7
        elif final_tabel_a >= 8 and final_tabel_b == 5:  # 85
            Tabel_c = 7
        elif final_tabel_a >= 8 and final_tabel_b == 6:  # 86
            Tabel_c = 7
        elif final_tabel_a >= 8 and final_tabel_b >= 7:  # 87
            Tabel_c = 7
        else:
            Tabel_c = 0
        
        # Tabel c
        if Tabel_c == 1 or Tabel_c == 2:
            final_score ="Acceptable posture"
        elif Tabel_c == 3 or Tabel_c == 4:
            final_score = "Further investigation, change may be needed"
        elif Tabel_c == 5 or Tabel_c == 6:
            final_score = "Further investigation, change soon"
        elif Tabel_c == 7:
            final_score = "Investigate and implement change"
        else:
            final_score = "Invalid value for Tabel_c"
        
        # Simpan data sudut-sudut
        # if midear[0] < frame.shape[1] // 2:
        if image_view == "Right":
            angles_data['shoulder_angle'] = angle_deg_2
            angles_data['elbow_angle'] = angle_deg_4
            angles_data['wrist_angle'] = abs(angle_deg_6-180)

        elif image_view == "Left":
            angles_data['shoulder_angle'] = angle_deg_1
            angles_data['elbow_angle'] = angle_deg_3
            angles_data['wrist_angle'] = abs(angle_deg_5-180)
        else:
            print("No valid image view selected")
        
        #Neck and Trunk
        angles_data['neck_angle'] = angle_deg_8
        angles_data['trunk_angle'] = abs(angle_deg_9 - 180)
        # Memformat angka dengan dua desimal sebelum mengirimkannya ke template
        # Function to clean up the string data
        def clean_data(data):
            # Replace specific unwanted strings or characters
            return [item.replace('[', '').replace(']', '').replace("'", '').strip() for item in data]

        # Process the adjustments data here and clean it
        upper_arm_position_str = clean_data(adjustments.get('upper_arm_position', []))
        lower_arm_position_str = clean_data(adjustments.get('lower_arm_position', []))
        wrist_position_str = clean_data(adjustments.get('wrist_position', []))
        wrist_twist_str = clean_data(adjustments.get('wrist_twist', []))
        neck_position_str = clean_data(adjustments.get('neck_position', []))
        trunk_position_str = clean_data(adjustments.get('trunk_position', []))
        legs_position_str = clean_data(adjustments.get('legs_position', []))
        add_muscle_a_str = clean_data(adjustments.get('add_muscle_a', []))
        add_muscle_b_str = clean_data(adjustments.get('add_muscle_b', []))
        add_force_a_str = clean_data(adjustments.get('add_force_a', []))
        add_force_b_str = clean_data(adjustments.get('add_force_b', []))

        # Mapping from description to numeric value
        value_map = {
            'Shoulder is raised': 1,
            'Upper arm is abducted': 1,
            'Arm is supported or person is leaning': -1,
            'Arm working across midline or out to side': 1,
            'Wrist bent from midline': 1,
            'Twisted in mid-range': 1,
            'At or near end of range': 2,
            'Neck is twisted': 1,
            'Neck is side bending': 1,
            'Trunk is twisted': 1,
            'Trunk is side bending': 1,
            'Legs and feet are supported': 1,
            'Not supported': 2,
            'If posture mainly static (i.e. held>10 minutes), Or if action repeated occurs 4x per minute:': 1,
            'If load < 4.4 lbs. (intermittent)': 0,
            'If load 4.4 to 22 lbs. (intermittent)': 1,
            'If load 4.4 to 22 lbs. (static or repeated)': 2,
            'If more than 22 lbs. or repeated or shocks': 3
        }

        def convert_to_numeric(adjustment_list):
            return [value_map.get(item, 0) for item in adjustment_list]

        # Convert adjustments to numeric values and sum them
        upper_arm_position_int = sum(convert_to_numeric(upper_arm_position_str))
        lower_arm_position_int = sum(convert_to_numeric(lower_arm_position_str))
        wrist_position_int = sum(convert_to_numeric(wrist_position_str))
        wrist_twist_int = sum(convert_to_numeric(wrist_twist_str))
        neck_position_int = sum(convert_to_numeric(neck_position_str))
        trunk_position_int = sum(convert_to_numeric(trunk_position_str))
        legs_position_int = sum(convert_to_numeric(legs_position_str))
        add_muscle_a_int = sum(convert_to_numeric(add_muscle_a_str))
        add_muscle_b_int = sum(convert_to_numeric(add_muscle_b_str))
        add_force_a_int = sum(convert_to_numeric(add_force_a_str))
        add_force_b_int = sum(convert_to_numeric(add_force_b_str))

        # Simpan data total score dalam score_data
        score_data = {
            'upper_arm_score': Upper_Arm_Score,
            'lower_arm_score': Lower_Arm_Score,
            'wrist_score': Wrist_Score,
            'neck_score': Neck_Score,
            'trunk_score': Trunk_Score,
            'wrist_twist': wrist_twist,
            'legs_position': legs_position,
            'add_muscle_a': add_muscle_a,
            'add_muscle_b': add_muscle_b,
            'add_force_a': add_force_a,
            'add_force_b': add_force_b,
            'upper_arm_position': upper_arm_position,
            'lower_arm_position': lower_arm_position,
            'wrist_position': wrist_position,
            'neck_position': neck_position,
            'trunk_position': trunk_position,
            'total_upper_arm_score': total_upper_arm_score,
            'total_lower_arm_score': total_lower_arm_score,
            'total_wrist_score': total_wrist_score,
            'total_neck_score': total_neck_score,
            'total_trunk_score': total_trunk_score,
            'Tabel_a': Tabel_a,
            'Tabel_b': Tabel_b,
            'Tabel_c': Tabel_c,
            'final_score': final_score,
            'final_tabel_a': final_tabel_a,
            'final_tabel_b': final_tabel_b,
            'image_view': image_view,
            'upper_arm_position_int': upper_arm_position_int,
            'lower_arm_position_int': lower_arm_position_int,
            'wrist_position_int': wrist_position_int,
            'wrist_twist_int': wrist_twist_int,
            'neck_position_int': neck_position_int,
            'trunk_position_int': trunk_position_int,
            'legs_position_int': legs_position_int,
            'add_muscle_a_int': add_muscle_a_int,
            'add_muscle_b_int': add_muscle_b_int,
            'add_force_a_int': add_force_a_int,
            'add_force_b_int': add_force_b_int,
            'upper_arm_position_str': upper_arm_position_str,
            'lower_arm_position_str': lower_arm_position_str,
            'wrist_position_str': wrist_position_str,
            'wrist_twist_str': wrist_twist_str,
            'neck_position_str': neck_position_str,
            'trunk_position_str': trunk_position_str,
            'legs_position_str': legs_position_str,
            'add_muscle_a_str': add_muscle_a_str,
            'add_muscle_b_str': add_muscle_b_str,
            'add_force_a_str': add_force_a_str,
            'add_force_b_str': add_force_b_str
        }
        
        desc_data = {
            'worker_name': worker_name,
            'job_location': job_location,
            'job_name': job_name,
        }

        # Simpan frame yang sudah diolah
        # Simpan gambar yang telah diproses
    processed_filename = 'processed_' + latest_file
    processed_path = os.path.join(app.config['PROCESSED_FOLDER'], processed_filename)
    cv2.imwrite(processed_path, rgb_frame)
    
    # Simpan path gambar yang diproses ke dalam session
    session['processed_filename'] = processed_filename

    # Menyimpan hasil gambar
    result_filename = 'result_' + latest_file  # Membuat nama file hasil dengan awalan 'result_'
    result_path = os.path.join(app.config['RESULTS_FOLDER'], result_filename)
    cv2.imwrite(os.path.join('static/results/', result_filename), frame)
    # Simpan path gambar yang diproses ke dalam session
    session['result_filename'] = result_filename

    # Memformat angka dengan dua desimal sebelum mengirimkannya ke template
    for key, value in angles_data.items():
        angles_data[key] = "{:.2f}".format(float(value))
    
    session['angles_data_s'] = angles_data
    session['score_data_s'] = score_data
    session['desc_data_s'] = desc_data

    return render_template('result.html', image_path='static/results/' + result_filename, angles_data=angles_data, score_data=score_data, desc_data=desc_data)

@app.route('/download_docx', methods=['POST'])
def download_docx():
    # Ubah sesuai dengan format data yang Anda kirim
    angles_data = session.get('angles_data_s')
    score_data = session.get('score_data_s')
    desc_data = session.get('desc_data_s')
    result_filename = session.get('result_filename')
    
    if not result_filename:
        flash('No processed image available for embedding.')
        return redirect(url_for('index'))

    # Buat struktur data yang sesuai
    data = {
        'angles_data': angles_data,
        'score_data': score_data,
        'desc_data': desc_data,
        'temp_img_path': os.path.join(app.config['RESULTS_FOLDER'], result_filename) # Gabungkan dengan path lengkap
    }
    
    # Generate document
    docx_path = generate_document(data)

    # Serve dokumen sebagai file yang dapat diunduh
    response = make_response(open(docx_path, 'rb').read())
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    response.headers['Content-Disposition'] = 'attachment; filename=result.docx'

    return response

def generate_document(data):
    # Ambil path file gambar dari data
    temp_img_path = data.get('temp_img_path')
    
    # Pastikan file gambar ada sebelum digunakan
    if not temp_img_path or not os.path.exists(temp_img_path):
        raise FileNotFoundError(f"File '{temp_img_path}' not found or invalid.")

    # Baca template dokumen
    template_path = 'templates/evaluation_sheet - Copy.docx'
    document = Document(template_path)

    # Ganti placeholder dengan data yang sesuai
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data['angles_data'].items():
                    placeholder = "{{" + key + "}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))

                        # Ubah ukuran font
                        for paragraph in cell.paragraphs:
                            run = paragraph.runs[0]
                            run.font.size = Pt(9)  # Ubah ke ukuran font yang diinginkan

                            # Ubah jenis font
                            run.font.name = 'Arial'  # Ubah ke jenis font yang diinginkan

                            # Atur paragraf menjadi pusat (centered)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                for key, value in data['score_data'].items():
                    placeholder = "{{" + key + "}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))

                        # Ubah ukuran font
                        for paragraph in cell.paragraphs:
                            run = paragraph.runs[0]
                            run.font.size = Pt(9)  # Ubah ke ukuran font yang diinginkan

                            # Ubah jenis font
                            run.font.name = 'Arial'  # Ubah ke jenis font yang diinginkan

                            # Atur paragraf menjadi pusat (centered)
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                for key, value in data['desc_data'].items():
                    placeholder = "{{" + key + "}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))

                        # Ubah ukuran font
                        for paragraph in cell.paragraphs:
                            run = paragraph.runs[0]
                            run.font.size = Pt(9)  # Ubah ke ukuran font yang diinginkan

                            # Ubah jenis font
                            run.font.name = 'Arial'  # Ubah ke jenis font yang diinginkan
                                
                # Ganti placeholder untuk gambar
                if '{{image_path}}' in cell.text:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if '{{image_path}}' in run.text:
                                run.text = run.text.replace('{{image_path}}', '')
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                # Sisipkan gambar
                                paragraph.add_run().add_picture(temp_img_path, width=Inches(3))

    # Simpan dokumen yang dihasilkan
    output_dir = 'output'
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    output_path = os.path.join(output_dir, 'result.docx')
    document.save(output_path)
    return output_path

if __name__ == '__main__':
    app.run(debug=True)
